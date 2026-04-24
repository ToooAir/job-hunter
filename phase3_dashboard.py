"""
Phase 3 — Apply Assistant (Streamlit Dashboard)
Run: streamlit run phase3_dashboard.py
"""

import hashlib
import json
from datetime import datetime, timezone, timedelta
from pathlib import Path
from urllib.parse import quote as url_quote

import pandas as pd
import pyperclip
import streamlit as st
import yaml

from utils.db import (
    init_db, upsert_job, update_status, set_follow_up, set_notes,
    add_interview_record, get_interview_records, delete_interview_record,
    get_company_applications,
)

# ── Page config ────────────────────────────────────────────────────────────────

st.set_page_config(layout="wide", page_title="Job Dashboard")

# ── i18n ───────────────────────────────────────────────────────────────────────

STRINGS: dict[str, dict[str, str]] = {
    "en": {
        # KPIs
        "kpi_pending":        "Pending Review 🎯",
        "kpi_week_applied":   "Applied (Last 7 Days) ✅",
        "kpi_in_interview":   "In Interview 📞",
        "kpi_offer":          "Offer 🎉",
        "kpi_followup":       "Follow-up Due 🔔",
        "kpi_ghosted":        "Ghosted 👻",
        "kpi_errors":         "Score Errors ❌",
        # Stats
        "stats_expander":     "📊 Analytics",
        "grade_dist":         "**Grade Distribution**",
        "lang_dist":          "**Language Requirement Distribution**",
        "funnel_title":       "**Application Funnel**",
        "source_eff":         "**Source Effectiveness**",
        "weekly_trend":       "**Weekly Application Trend (last 8 weeks)**",
        "daily_trend":        "**This Week — Applications by Day**",
        "no_scored_jobs":     "No scored jobs yet",
        "no_applications":    "No applications yet",
        "no_data":            "No data yet",
        # Funnel stage keys → display labels
        "stage_applied":      "Applied",
        "stage_interview_1":  "1st Interview",
        "stage_interview_2":  "2nd Interview",
        "stage_offer":        "Offer",
        # src_df column display names
        "col_source":         "Source",
        "col_total":          "Total",
        "col_grade_a":        "Grade A",
        "col_grade_b":        "Grade B",
        "col_applied":        "Applied",
        "col_interviews":     "Interviews",
        "col_grade_a_pct":    "A-Grade %",
        "col_reply_rate":     "Reply Rate",
        "col_count":          "Count",
        # Pipeline log
        "log_expander":       "🕐 Pipeline Schedule Log",
        "no_log":             "No pipeline run recorded yet. Will appear after first manual run or scheduled trigger.",
        "run_btn":            "▶️ Run Phase 1 + 2 Now",
        "refresh_btn":        "🔄 Refresh Log",
        "retry_errors_btn":   "⚠️ Retry All Errors",
        "retry_errors_spin":  "Resetting & scoring error jobs…",
        "retry_errors_none":  "No error-status jobs found.",
        "run_starting":       "Starting pipeline (background)…",
        "run_started":        "Pipeline started. Refresh the page in a moment to see results.",
        "log_caption":        "logs/pipeline.log — {size:.1f} KB, showing last 100 lines",
        # Filters
        "filter_expander":    "Filters",
        "filter_grade":       "Grade",
        "filter_lang":        "Language Req",
        "filter_source":      "Source",
        "filter_status":      "Status",
        "no_jobs":            "No jobs match the current filters.",
        # LinkedIn
        "linkedin_search":    "**LinkedIn Search**",
        # Search targets manager
        "mgr_expander":       "⚙️ Manage Search Targets",
        "tab_kw":             "Keywords",
        "tab_gh":             "Greenhouse",
        "tab_lv":             "Lever",
        "kw_hint":            "Keywords shared by Arbeitnow and Bundesagentur (one per line)",
        "ba_kw_label":        "Bundesagentur Keywords",
        "save_kw_btn":        "💾 Save Keywords",
        "save_kw_ok":         "✅ Keywords saved",
        "gh_hint":            "One slug per line (no quotes). 404 slugs are skipped at runtime.",
        "gh_label":           "Greenhouse Company Slugs",
        "save_gh_btn":        "💾 Save Greenhouse List",
        "save_gh_ok":         "✅ Greenhouse list saved",
        "lv_hint":            "One slug per line (no quotes). 404 slugs are skipped at runtime.",
        "lv_label":           "Lever Company Slugs",
        "save_lv_btn":        "💾 Save Lever List",
        "save_lv_ok":         "✅ Lever list saved",
        # Manual add
        "manual_expander":    "Add Job Manually",
        "company_field":      "Company Name *",
        "title_field":        "Job Title *",
        "url_field":          "Job URL *",
        "location_field":     "Location",
        "source_field":       "Source",
        "jd_field":           "Job Description (paste full JD) *",
        "manual_submit":      "Add to Database",
        "manual_required":    "Please fill in all required fields (marked *)",
        "manual_added":       "✅ Added to database. Run `python phase2_scorer.py` then refresh.",
        "manual_exists":      "⚠️ This URL already exists in the database.",
        # Right column
        "select_prompt":      "← Select a job from the left",
        "not_found":          "Job not found, please refresh the page.",
        "translated_flag":    "　🔤 JD auto-translated (original: German)",
        # Status labels (for duplicate warning)
        "status_applied":     "Applied",
        "status_interview_1": "1st Interview",
        "status_interview_2": "2nd Interview",
        "status_offer":       "Offer",
        "status_rejected":    "Rejected",
        "status_ghosted":     "Ghosted",
        # Duplicate warning
        "dup_warning":        "⚠️ You have previously applied to other roles at **{company}**: {roles}",
        # Visa
        "visa_eu_only":       "🚫 EU Only — analyze Chancenkarte compatibility before applying",
        "visa_sponsored":     "✅ Company offers visa sponsorship — Chancenkarte friendly",
        "visa_open":          "🟢 No explicit restriction",
        "visa_unclear":       "⚪ No visa requirement mentioned",
        "visa_expander_pfx":  "🛂 Visa Compatibility: ",
        "visa_warning_text":  "JD contains work permit restriction language, but Chancenkarte may not be excluded. Click below for deep analysis.",
        "visa_sponsored_txt": "Company actively offers visa assistance — safe to apply.",
        "visa_unclear_txt":   "Click below to analyze JD compatibility with Chancenkarte.",
        "visa_analyze_btn":   "🔍 Deep Analysis: Chancenkarte Compatibility",
        "visa_reanalyze_btn": "🔄 Re-analyze",
        "visa_analyzing":     "Analyzing…",
        # JD viewer
        "jd_expander":        "📄 Full Job Description",
        "jd_tab_en":          "English (translated)",
        "jd_tab_de":          "German (original)",
        "jd_empty":           "(No JD content)",
        # Salary
        "salary_prefix":      "💰 Salary Estimate",
        "salary_jd_shown":    "　(JD states: {salary})",
        "salary_not_gen":     "Salary estimate not yet generated.",
        "salary_jd_cap":      "JD-listed salary: {salary}",
        "salary_gen_btn":          "💰 Generate Salary Estimate",
        "salary_regen_btn":        "🔄 Re-estimate",
        "salary_estimating":       "Estimating…",
        "levels_cache_fresh":      "Levels.fyi ref data: {location} ({count} layers), {days}d ago",
        "levels_cache_stale":      "Levels.fyi ref data: {location} ({count} layers), {days}d ago (stale)",
        "levels_cache_none":       "No Levels.fyi cache — will fetch on next estimate",
        "levels_refresh_btn":      "🔄 Refresh Levels data",
        "levels_refreshing":       "Fetching Levels.fyi data…",
        "levels_refresh_all_btn":  "🔄 Refresh All Levels Data",
        "levels_refresh_all_done": "✓ Levels.fyi data refreshed ({n} slots)",
        "levels_refresh_all_spin": "Fetching Levels.fyi data for all roles…",
        # Company research
        "research_expander":  "🔍 Company Research",
        "research_not_gen":   "Company research not yet generated.",
        "research_gen_btn":   "🔍 Research Company",
        "research_regen_btn": "🔄 Regenerate",
        "research_spinner":   "Researching {company}…",
        "kununu_btn":         "🌐 Kununu Reviews",
        # Reasons
        "reasons_title":      "**📌 Scoring Reasons**",
        "not_scored_cap":     "(Not yet scored)",
        # Cover letter
        "cl_title":           "**✉️ Cover Letter (editable)**",
        "cl_word_count":      "Word count: {n} / recommended 200–400",
        "cl_too_long":        "⚠️ Cover letter exceeds 400 words — consider trimming before copying",
        "tone_label":         "Tone",
        "tone_formal":        "🏛️ Formal — Professional corporate style",
        "tone_startup":       "🚀 Startup — Direct and personable",
        "tone_concise":       "✂️ Concise — Under 200 words, precise",
        "regen_cl_btn":       "🔄 Regenerate Cover Letter",
        "regen_cl_spinner":   "Regenerating with {tone} tone…",
        "download_docx":      "📄 Download .docx",
        "download_pdf":       "📄 Download .pdf",
        # Action buttons
        "open_job_btn":       "🌐 Open Job Page",
        "copy_cl_btn":        "📋 Copy Cover Letter",
        "apply_btn":          "✅ Applied",
        "skip_btn":           "⏭️ Skip",
        "rescore_btn":        "🔄 Re-score",
        "rescore_spinner":    "Scoring…",
        "copied_ok":          "Copied ✓",
        "copy_docker_msg":    "No clipboard in Docker — please select and copy the cover letter text manually.",
        "iv1_btn":            "📞 Interview Invite",
        "iv1_spinner":        "Generating interview brief…",
        "reject_btn":         "❌ Rejected",
        "iv2_btn":            "💻 2nd Interview",
        "offer_btn":          "🎉 Received Offer",
        "expired_warning":    "⏰ This job has passed its expiry date (expires_at has passed)",
        # Interview brief
        "brief_title":        "**🎯 Interview Brief**",
        "brief_regen_btn":    "🔄 Regenerate",
        "brief_not_gen":      "(Interview brief not yet generated — click 🔄 Regenerate to create one)",
        "brief_gen_spinner":  "Generating…",
        "brief_download":     "⬇️ Download Markdown",
        # Interview records
        "records_title":      "**📋 Interview Records**",
        "questions_label":    "**Questions Asked**",
        "impressions_label":  "**Impressions**",
        "del_record_btn":     "🗑️ Delete this record",
        "add_record_exp":     "➕ Add Interview Record",
        "round_label":        "Round",
        "ir_date_label":      "Interview Date",
        "ir_format_label":    "Format",
        "interviewer_label":  "Interviewer (name / title)",
        "questions_field":    "Questions asked",
        "rating_label":       "Self-rating",
        "rating_help":        "1 = Poor, 5 = Great",
        "impressions_field":  "Impressions / notes",
        "save_record_btn":    "💾 Save Record",
        # Round / format labels
        "round_iv1":          "Round 1",
        "round_iv2":          "Round 2",
        "round_other":        "Other",
        "fmt_phone":          "☎️ Phone",
        "fmt_video":          "💻 Video",
        "fmt_onsite":         "🏢 On-site",
        "fmt_technical":      "⌨️ Technical",
        # Notes
        "notes_title":        "**📝 Notes**",
        "notes_placeholder":  "Interview impressions, contact info, salary range…",
        "save_notes_btn":     "Save Notes",
        "notes_saved":        "Notes saved ✓",
        # Follow-up
        "followup_title":     "**🔔 Follow-up Reminder**",
        "followup_save_btn":  "Save",
        "followup_overdue":   "⚠️ Follow-up date reached ({date}) — remember to send a follow-up!",
        "followup_days_left": "Follow-up in {n} day(s) ({date})",
        # Contract
        "contract_permanent": "✅ Permanent",
        "contract_contract":  "⏳ Contract",
        "contract_freelance": "🔧 Freelance",
        "contract_unknown":   "— Unknown",
        # Location filter
        "filter_location":    "Location Search",
        "filter_location_ph": "e.g. Hamburg",
        "filter_remote":      "Also include Remote jobs",
        # Lang toggle
        "lang_toggle_label":  "🌐 Language",
    },
    "zh": {
        # KPIs
        "kpi_pending":        "待審閱 🎯",
        "kpi_week_applied":   "過去 7 天投遞 ✅",
        "kpi_in_interview":   "面試中 📞",
        "kpi_offer":          "Offer 🎉",
        "kpi_followup":       "待跟進 🔔",
        "kpi_ghosted":        "無聲卡 👻",
        "kpi_errors":         "評分失敗 ❌",
        # Stats
        "stats_expander":     "📊 統計分析",
        "grade_dist":         "**等級分布**",
        "lang_dist":          "**語言要求分布**",
        "funnel_title":       "**應聘漏斗**",
        "source_eff":         "**來源效益**",
        "weekly_trend":       "**每週投遞趨勢（近 8 週）**",
        "daily_trend":        "**本週每日投遞數**",
        "no_scored_jobs":     "尚無已評分職缺",
        "no_applications":    "尚無投遞記錄",
        "no_data":            "尚無資料",
        # Funnel stages
        "stage_applied":      "已投遞",
        "stage_interview_1":  "一面邀約",
        "stage_interview_2":  "二面邀約",
        "stage_offer":        "Offer",
        # src_df columns
        "col_source":         "來源",
        "col_total":          "總筆數",
        "col_grade_a":        "A 級",
        "col_grade_b":        "B 級",
        "col_applied":        "已投遞",
        "col_interviews":     "進面試",
        "col_grade_a_pct":    "A 級率",
        "col_reply_rate":     "回覆率",
        "col_count":          "筆數",
        # Pipeline log
        "log_expander":       "🕐 Pipeline 排程記錄",
        "no_log":             "尚無 pipeline 執行記錄。首次手動執行或等排程觸發後會出現。",
        "run_btn":            "▶️ 立即執行 Phase 1 + 2",
        "refresh_btn":        "🔄 重新整理記錄",
        "retry_errors_btn":   "⚠️ 重試所有失敗",
        "retry_errors_spin":  "重設並評分失敗職缺中…",
        "retry_errors_none":  "目前沒有 error 狀態的職缺。",
        "run_starting":       "啟動 pipeline（背景執行）…",
        "run_started":        "已啟動，稍後重新整理頁面查看結果。",
        "log_caption":        "logs/pipeline.log — {size:.1f} KB，顯示最後 100 行",
        # Filters
        "filter_expander":    "篩選條件",
        "filter_grade":       "等級",
        "filter_lang":        "語言要求",
        "filter_source":      "來源",
        "filter_status":      "狀態",
        "no_jobs":            "沒有符合條件的職缺。",
        # LinkedIn
        "linkedin_search":    "**LinkedIn 搜尋**",
        # Search targets manager
        "mgr_expander":       "⚙️ 管理搜尋條件",
        "tab_kw":             "關鍵字",
        "tab_gh":             "Greenhouse",
        "tab_lv":             "Lever",
        "kw_hint":            "Arbeitnow 與 Bundesagentur 共用這份關鍵字清單（一行一個）",
        "ba_kw_label":        "Bundesagentur 關鍵字",
        "save_kw_btn":        "💾 儲存關鍵字",
        "save_kw_ok":         "✅ 關鍵字已儲存",
        "gh_hint":            "一行一個 slug（不需加引號）。404 的 slug 執行時自動跳過。",
        "gh_label":           "Greenhouse 公司 slug",
        "save_gh_btn":        "💾 儲存 Greenhouse 清單",
        "save_gh_ok":         "✅ Greenhouse 清單已儲存",
        "lv_hint":            "一行一個 slug（不需加引號）。404 的 slug 執行時自動跳過。",
        "lv_label":           "Lever 公司 slug",
        "save_lv_btn":        "💾 儲存 Lever 清單",
        "save_lv_ok":         "✅ Lever 清單已儲存",
        # Manual add
        "manual_expander":    "手動新增職缺",
        "company_field":      "公司名稱 *",
        "title_field":        "職位名稱 *",
        "url_field":          "職缺 URL *",
        "location_field":     "地點",
        "source_field":       "來源",
        "jd_field":           "職缺描述（貼入 JD 全文）*",
        "manual_submit":      "加入資料庫",
        "manual_required":    "請填入所有必填欄位（標 * 者）",
        "manual_added":       "✅ 已加入資料庫。請執行 `python phase2_scorer.py` 後重新整理頁面。",
        "manual_exists":      "⚠️ 此 URL 已存在於資料庫中。",
        # Right column
        "select_prompt":      "← 從左側選擇一筆職缺",
        "not_found":          "找不到該職缺，請重新整理頁面。",
        "translated_flag":    "　🔤 JD 已自動翻譯（原文德語）",
        # Status labels
        "status_applied":     "已投遞",
        "status_interview_1": "一面",
        "status_interview_2": "二面",
        "status_offer":       "Offer",
        "status_rejected":    "已拒絕",
        "status_ghosted":     "無聲卡",
        # Duplicate warning
        "dup_warning":        "⚠️ 你曾投遞過 **{company}** 的其他職缺：{roles}",
        # Visa
        "visa_eu_only":       "🚫 EU Only — 投遞前建議深度分析 Chancenkarte 相容性",
        "visa_sponsored":     "✅ 公司提供簽證協助 — Chancenkarte 友善",
        "visa_open":          "🟢 無明確限制",
        "visa_unclear":       "⚪ 未提及簽證要求",
        "visa_expander_pfx":  "🛂 簽證相容性：",
        "visa_warning_text":  "JD 含有工作許可限制字句，但 Chancenkarte 不一定被排除。點下方按鈕做深度分析。",
        "visa_sponsored_txt": "公司主動提供簽證協助，可安心投遞。",
        "visa_unclear_txt":   "點下方按鈕分析 JD 對 Chancenkarte 的相容性。",
        "visa_analyze_btn":   "🔍 深度分析 Chancenkarte 相容性",
        "visa_reanalyze_btn": "🔄 重新分析",
        "visa_analyzing":     "分析中…",
        # JD viewer
        "jd_expander":        "📄 完整 JD 原文",
        "jd_tab_en":          "英文（已翻譯）",
        "jd_tab_de":          "德文（原文）",
        "jd_empty":           "（無 JD 內容）",
        # Salary
        "salary_prefix":      "💰 薪資估計",
        "salary_jd_shown":    "　（JD 標示：{salary}）",
        "salary_not_gen":     "尚未生成薪資估計。",
        "salary_jd_cap":      "JD 標示薪資：{salary}",
        "salary_gen_btn":          "💰 生成薪資估計",
        "salary_regen_btn":        "🔄 重新估計",
        "salary_estimating":       "估算中…",
        "levels_cache_fresh":      "Levels.fyi 參考資料：{location}（{count} 層），{days} 天前更新",
        "levels_cache_stale":      "Levels.fyi 參考資料：{location}（{count} 層），{days} 天前更新（已過期）",
        "levels_cache_none":       "尚無 Levels.fyi 快取，下次估計時將自動抓取",
        "levels_refresh_btn":      "🔄 刷新 Levels 資料",
        "levels_refreshing":       "正在抓取 Levels.fyi 資料…",
        "levels_refresh_all_btn":  "🔄 刷新全部 Levels 資料",
        "levels_refresh_all_done": "✓ Levels.fyi 資料已更新（{n} 個 slot）",
        "levels_refresh_all_spin": "正在抓取所有職位類型的 Levels.fyi 資料…",
        # Company research
        "research_expander":  "🔍 公司研究",
        "research_not_gen":   "尚未生成公司研究報告。",
        "research_gen_btn":   "🔍 生成公司研究",
        "research_regen_btn": "🔄 重新生成",
        "research_spinner":   "正在研究 {company}…",
        "kununu_btn":         "🌐 Kununu 評價",
        # Reasons
        "reasons_title":      "**📌 評分理由**",
        "not_scored_cap":     "（尚未評分）",
        # Cover letter
        "cl_title":           "**✉️ Cover Letter（可直接編輯）**",
        "cl_word_count":      "字數：{n} / 建議 200–400 字",
        "cl_too_long":        "⚠️ Cover Letter 超過 400 字，建議壓縮後再複製",
        "tone_label":         "語氣",
        "tone_formal":        "🏛️ Formal — 正式企業風格",
        "tone_startup":       "🚀 Startup — 直接有個性",
        "tone_concise":       "✂️ Concise — 200 字以內，精準",
        "regen_cl_btn":       "🔄 重新生成 Cover Letter",
        "regen_cl_spinner":   "以 {tone} 語氣重新生成中…",
        "download_docx":      "📄 下載 .docx",
        "download_pdf":       "📄 下載 .pdf",
        # Action buttons
        "open_job_btn":       "🌐 開啟職缺頁面",
        "copy_cl_btn":        "📋 複製 Cover Letter",
        "apply_btn":          "✅ 已投遞",
        "skip_btn":           "⏭️ 略過",
        "rescore_btn":        "🔄 重新評分",
        "rescore_spinner":    "評分中…",
        "copied_ok":          "已複製 ✓",
        "copy_docker_msg":    "Docker 環境無剪貼簿，請手動選取 Cover Letter 文字後複製。",
        "iv1_btn":            "📞 面試邀約",
        "iv1_spinner":        "正在生成面試準備單…",
        "reject_btn":         "❌ 已拒絕",
        "iv2_btn":            "💻 二次面試",
        "offer_btn":          "🎉 收到 Offer",
        "expired_warning":    "⏰ 此職缺已超過有效期限（expires_at 已過）",
        # Interview brief
        "brief_title":        "**🎯 面試準備單**",
        "brief_regen_btn":    "🔄 重新生成",
        "brief_not_gen":      "（尚未生成面試準備單 — 點「🔄 重新生成」手動產生）",
        "brief_gen_spinner":  "生成中…",
        "brief_download":     "⬇️ 下載 Markdown",
        # Interview records
        "records_title":      "**📋 面試記錄**",
        "questions_label":    "**被問到的問題**",
        "impressions_label":  "**感想**",
        "del_record_btn":     "🗑️ 刪除此記錄",
        "add_record_exp":     "➕ 新增面試記錄",
        "round_label":        "輪次",
        "ir_date_label":      "面試日期",
        "ir_format_label":    "形式",
        "interviewer_label":  "面試官（姓名 / 職稱）",
        "questions_field":    "被問到的問題",
        "rating_label":       "自我感覺",
        "rating_help":        "1 = 很差，5 = 很好",
        "impressions_field":  "感想 / 注意事項",
        "save_record_btn":    "💾 儲存記錄",
        # Round / format labels
        "round_iv1":          "第一輪",
        "round_iv2":          "第二輪",
        "round_other":        "其他",
        "fmt_phone":          "☎️ 電話",
        "fmt_video":          "💻 視訊",
        "fmt_onsite":         "🏢 現場",
        "fmt_technical":      "⌨️ 技術面試",
        # Notes
        "notes_title":        "**📝 備註**",
        "notes_placeholder":  "面試感想、聯絡人、薪資範圍…",
        "save_notes_btn":     "儲存備註",
        "notes_saved":        "備註已儲存 ✓",
        # Follow-up
        "followup_title":     "**🔔 跟進提醒**",
        "followup_save_btn":  "儲存",
        "followup_overdue":   "⚠️ 跟進日期已到（{date}）— 記得發跟進信！",
        "followup_days_left": "距跟進日期還有 {n} 天（{date}）",
        # Contract
        "contract_permanent": "✅ 正職",
        "contract_contract":  "⏳ 合約",
        "contract_freelance": "🔧 Freelance",
        "contract_unknown":   "— unknown",
        # Location filter
        "filter_location":    "地點搜尋",
        "filter_location_ph": "例：Hamburg",
        "filter_remote":      "也包含 Remote 職缺",
        # Lang toggle
        "lang_toggle_label":  "🌐 語言",
    },
}


def _lang() -> str:
    return st.session_state.get("lang", "zh")


def T(key: str) -> str:  # noqa: N802
    return STRINGS[_lang()].get(key, STRINGS["en"].get(key, key))


# ── Language toggle (sidebar) ─────────────────────────────────────────────────

_lang_map = {"中文": "zh", "English": "en"}
_lang_display = st.sidebar.radio(
    T("lang_toggle_label"),
    list(_lang_map.keys()),
    index=0 if _lang() == "zh" else 1,
    horizontal=True,
)
st.session_state["lang"] = _lang_map[_lang_display]

# ── Levels.fyi global refresh (sidebar) ───────────────────────────────────────

st.sidebar.divider()
if st.sidebar.button(T("levels_refresh_all_btn"), use_container_width=True):
    import os
    from dotenv import load_dotenv
    from utils.levels_scraper import (
        clear_cache, fetch_levels_by_slug,
        _ROLE_MAP, _ROLE_FALLBACK, HOME_COUNTRY, FALLBACK_CHAIN,
    )
    load_dotenv()

    _seen: dict[str, None] = dict.fromkeys(slug for _, slug in _ROLE_MAP)
    _seen.setdefault(_ROLE_FALLBACK, None)
    _all_role_slugs = list(_seen)
    _loc_chain      = FALLBACK_CHAIN.get(HOME_COUNTRY, [HOME_COUNTRY, "global"])

    with st.sidebar:
        with st.spinner(T("levels_refresh_all_spin")):
            _refreshed = 0
            for _rs in _all_role_slugs:
                for _ls in _loc_chain:
                    clear_cache(_rs, _ls)
                # fetch by slug directly — no title→slug mapping needed
                fetch_levels_by_slug(_rs, HOME_COUNTRY)
                _refreshed += len(_loc_chain)

    st.sidebar.success(T("levels_refresh_all_done").format(n=_refreshed))

# ── Helpers ────────────────────────────────────────────────────────────────────


def make_id(url: str) -> str:
    return hashlib.sha256(url.encode()).hexdigest()[:16]


def utcnow_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S")


def week_ago_iso() -> str:
    return (datetime.now(timezone.utc) - timedelta(days=7)).strftime("%Y-%m-%dT%H:%M:%S")


def get_conn():
    import os
    from dotenv import load_dotenv
    load_dotenv()
    db_path = os.getenv("DB_PATH", "./data/jobs.db")
    return init_db(db_path)


@st.cache_data(ttl=30)
def load_config():
    with open("config/search_targets.yaml", encoding="utf-8") as f:
        return yaml.safe_load(f)


def today_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d")


PIPELINE_STATUSES = ('applied', 'interview_1', 'interview_2', 'offer', 'rejected', 'ghosted')

# City alias expansion: handles English/German name variants and common misspellings
_GERMANY_PATTERNS = [
    # country-level markers (catches "Berlin, Germany", "Germany", "Deutschland", "bundesweit")
    "germany", "deutschland", "bundesweit",
    # major cities (English + German + Anglicised spellings)
    "hamburg", "berlin",
    "munich", "münchen", "muenchen",
    "cologne", "köln", "koeln",
    "frankfurt",
    "düsseldorf", "dusseldorf",
    "stuttgart",
    "nuremberg", "nürnberg", "nuernberg",
    "leipzig",
    "hannover", "hanover",
    "bremen",
    "dresden",
    "essen", "dortmund", "bochum",
    "karlsruhe", "mannheim", "heidelberg",
    "augsburg", "freiburg",
    "wiesbaden", "mainz", "bonn",
    "kiel", "rostock", "lübeck", "luebeck",
    "konstanz", "ulm", "regensburg",
]

_LOCATION_ALIASES: dict[str, list[str]] = {
    # Germany-wide search
    "germany":     _GERMANY_PATTERNS,
    "deutschland": _GERMANY_PATTERNS,
    # individual cities
    "hamburg":    ["hamburg"],
    "berlin":     ["berlin"],
    "munich":     ["munich", "münchen", "muenchen"],
    "cologne":    ["cologne", "köln", "koeln"],
    "frankfurt":  ["frankfurt"],
    "dusseldorf": ["dusseldorf", "düsseldorf"],
    "stuttgart":  ["stuttgart"],
    "nuremberg":  ["nuremberg", "nürnberg"],
    "leipzig":    ["leipzig"],
    "hannover":   ["hannover", "hanover"],
    "bremen":     ["bremen"],
    "dresden":    ["dresden"],
}

def _location_patterns(kw: str) -> list[str]:
    return _LOCATION_ALIASES.get(kw.lower().strip(), [kw.lower().strip()])


def fetch_kpis(conn) -> dict:
    rows = conn.execute(f"""
        SELECT
            SUM(CASE WHEN fit_grade IN ('A','B') AND status='scored'         THEN 1 ELSE 0 END) AS pending,
            SUM(CASE WHEN status IN {PIPELINE_STATUSES}                       THEN 1 ELSE 0 END) AS total_applied,
            SUM(CASE WHEN status IN {PIPELINE_STATUSES} AND applied_at >= ?   THEN 1 ELSE 0 END) AS week_applied,
            SUM(CASE WHEN status IN ('interview_1','interview_2','offer')      THEN 1 ELSE 0 END) AS in_interview,
            SUM(CASE WHEN status='offer'                                       THEN 1 ELSE 0 END) AS offers,
            SUM(CASE WHEN status IN {PIPELINE_STATUSES} AND follow_up_at <= ? AND follow_up_at IS NOT NULL
                                                                               THEN 1 ELSE 0 END) AS needs_followup,
            SUM(CASE WHEN status='ghosted'                                      THEN 1 ELSE 0 END) AS ghosted,
            SUM(CASE WHEN status='error'                                        THEN 1 ELSE 0 END) AS errors
        FROM jobs
    """, (week_ago_iso(), today_iso())).fetchone()
    return {
        "pending":        rows["pending"]        or 0,
        "total_applied":  rows["total_applied"]  or 0,
        "week_applied":   rows["week_applied"]   or 0,
        "in_interview":   rows["in_interview"]   or 0,
        "offers":         rows["offers"]         or 0,
        "needs_followup": rows["needs_followup"] or 0,
        "ghosted":        rows["ghosted"]        or 0,
        "errors":         rows["errors"]         or 0,
    }


def fetch_jobs(conn, grades, langs, sources, statuses,
               location_kw: str = "", include_remote: bool = False) -> pd.DataFrame:
    if not (grades and langs and sources and statuses):
        return pd.DataFrame()

    def placeholders(lst):
        return ",".join("?" * len(lst))

    # Build optional location clause
    location_clause = ""
    location_params: list = []
    kw = location_kw.strip()
    if kw or include_remote:
        parts = []
        if kw:
            for pattern in _location_patterns(kw):
                parts.append("LOWER(location) LIKE ?")
                location_params.append(f"%{pattern}%")
        if include_remote:
            parts.extend([
                "LOWER(location) LIKE '%remote%'",
                "LOWER(location) LIKE '%home office%'",
                "LOWER(location) LIKE '%homeoffice%'",
            ])
        location_clause = f"AND ({' OR '.join(parts)})"

    # error / un-scored jobs have no fit_grade — bypass that filter for them
    sql = f"""
        SELECT id, title, company, location,
               fit_grade, match_score, jd_language_req, source, status
        FROM jobs
        WHERE (fit_grade IN ({placeholders(grades)}) OR fit_grade IS NULL)
          AND (jd_language_req IN ({placeholders(langs)}) OR jd_language_req IS NULL)
          AND source           IN ({placeholders(sources)})
          AND status           IN ({placeholders(statuses)})
          {location_clause}
        ORDER BY
            CASE fit_grade WHEN 'A' THEN 1 WHEN 'B' THEN 2 ELSE 3 END,
            match_score DESC
    """
    params = grades + langs + sources + statuses + location_params
    rows = conn.execute(sql, params).fetchall()
    return pd.DataFrame([dict(r) for r in rows])


def fetch_job_detail(conn, job_id: str) -> dict | None:
    row = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    return dict(row) if row else None


@st.cache_data(ttl=60)
def fetch_stats(_conn) -> dict:
    """Return all data needed for the stats expander (language-neutral column names)."""
    # Grade distribution
    grade_rows = _conn.execute("""
        SELECT fit_grade, COUNT(*) AS cnt
        FROM jobs
        WHERE status IN ('scored','applied','skipped') AND fit_grade IS NOT NULL
        GROUP BY fit_grade
    """).fetchall()
    grade_df = pd.DataFrame([dict(r) for r in grade_rows]).set_index("fit_grade") if grade_rows else pd.DataFrame()

    # Language breakdown
    lang_rows = _conn.execute("""
        SELECT COALESCE(jd_language_req, 'unknown') AS lang, COUNT(*) AS cnt
        FROM jobs
        WHERE status IN ('scored','applied','skipped')
        GROUP BY lang
    """).fetchall()
    lang_df = pd.DataFrame([dict(r) for r in lang_rows]).set_index("lang") if lang_rows else pd.DataFrame()

    # Source effectiveness (keep English column names — translated at render time)
    src_rows = _conn.execute(f"""
        SELECT
            source,
            COUNT(*)                                                                        AS total,
            SUM(CASE WHEN fit_grade='A'                                  THEN 1 ELSE 0 END) AS grade_a,
            SUM(CASE WHEN fit_grade='B'                                  THEN 1 ELSE 0 END) AS grade_b,
            SUM(CASE WHEN status IN {PIPELINE_STATUSES}                  THEN 1 ELSE 0 END) AS applied,
            SUM(CASE WHEN status IN ('interview_1','interview_2','offer') THEN 1 ELSE 0 END) AS interviews
        FROM jobs
        WHERE status NOT IN ('un-scored','error')
        GROUP BY source
        ORDER BY grade_a DESC, total DESC
    """).fetchall()
    if src_rows:
        src_df = pd.DataFrame([dict(r) for r in src_rows])
        src_df["grade_a_pct"] = src_df.apply(
            lambda r: f"{r['grade_a'] / r['total'] * 100:.0f}%" if r["total"] else "—", axis=1
        )
        src_df["reply_rate"] = src_df.apply(
            lambda r: f"{r['interviews'] / r['applied'] * 100:.0f}%" if r["applied"] else "—", axis=1
        )
    else:
        src_df = pd.DataFrame()

    # Application funnel (English stage keys as index)
    f = _conn.execute(f"""
        SELECT
            SUM(CASE WHEN status IN {PIPELINE_STATUSES}                  THEN 1 ELSE 0 END) AS applied,
            SUM(CASE WHEN status IN ('interview_1','interview_2','offer') THEN 1 ELSE 0 END) AS interview_1,
            SUM(CASE WHEN status IN ('interview_2','offer')               THEN 1 ELSE 0 END) AS interview_2,
            SUM(CASE WHEN status='offer'                                  THEN 1 ELSE 0 END) AS offer
        FROM jobs
    """).fetchone()
    funnel_df = pd.DataFrame({
        "stage": ["applied", "interview_1", "interview_2", "offer"],
        "count": [f["applied"] or 0, f["interview_1"] or 0,
                  f["interview_2"] or 0, f["offer"] or 0],
    }).set_index("stage")

    # Weekly application trend (last 8 weeks)
    week_rows = _conn.execute(f"""
        SELECT strftime('%Y-W%W', applied_at) AS week, COUNT(*) AS cnt
        FROM jobs
        WHERE status IN {PIPELINE_STATUSES} AND applied_at IS NOT NULL
        GROUP BY week
        ORDER BY week DESC
        LIMIT 8
    """).fetchall()
    week_df = pd.DataFrame([dict(r) for r in week_rows]).set_index("week").sort_index() if week_rows else pd.DataFrame()

    # Daily breakdown — current Mon → today
    today = datetime.now(timezone.utc).date()
    week_start = today - timedelta(days=today.weekday())  # Monday
    daily_rows = _conn.execute(f"""
        SELECT strftime('%Y-%m-%d', applied_at) AS day, COUNT(*) AS cnt
        FROM jobs
        WHERE status IN {PIPELINE_STATUSES}
          AND applied_at IS NOT NULL
          AND strftime('%Y-%m-%d', applied_at) >= ?
          AND strftime('%Y-%m-%d', applied_at) <= ?
        GROUP BY day
        ORDER BY day
    """, (week_start.isoformat(), today.isoformat())).fetchall()
    # Ensure every day Mon→today appears (fill 0 for missing days)
    daily_counts = {r["day"]: r["cnt"] for r in daily_rows}
    days = []
    d = week_start
    while d <= today:
        days.append({"day": d.isoformat(), "cnt": daily_counts.get(d.isoformat(), 0)})
        d += timedelta(days=1)
    daily_df = pd.DataFrame(days).set_index("day")

    return {
        "grade_df":  grade_df,
        "lang_df":   lang_df,
        "src_df":    src_df,
        "funnel_df": funnel_df,
        "week_df":   week_df,
        "daily_df":  daily_df,
    }


GRADE_ICON = {"A": "🟢", "B": "🟡", "C": "🔴"}

# ── App ────────────────────────────────────────────────────────────────────────

conn = get_conn()
config = load_config()

# ── KPI row ────────────────────────────────────────────────────────────────────

kpis = fetch_kpis(conn)
k1, k2, k3, k4, k5, k6, k7 = st.columns(7)
k1.metric(T("kpi_pending"),       kpis["pending"])
k2.metric(T("kpi_week_applied"),  kpis["week_applied"])
k3.metric(T("kpi_in_interview"),  kpis["in_interview"])
k4.metric(T("kpi_offer"),         kpis["offers"])
k5.metric(T("kpi_followup"),      kpis["needs_followup"])
k6.metric(T("kpi_ghosted"),       kpis["ghosted"])
k7.metric(T("kpi_errors"),        kpis["errors"])

with st.expander(T("stats_expander"), expanded=False):
    stats = fetch_stats(conn)

    col_grade, col_lang = st.columns(2)

    with col_grade:
        st.markdown(T("grade_dist"))
        if not stats["grade_df"].empty:
            st.bar_chart(stats["grade_df"], y="cnt")
        else:
            st.caption(T("no_scored_jobs"))

    with col_lang:
        st.markdown(T("lang_dist"))
        if not stats["lang_df"].empty:
            st.bar_chart(stats["lang_df"], y="cnt")
        else:
            st.caption(T("no_scored_jobs"))

    st.markdown(T("funnel_title"))
    if not stats["funnel_df"].empty and stats["funnel_df"]["count"].sum() > 0:
        _f = stats["funnel_df"].copy()
        _f.index = [T(f"stage_{s}") for s in _f.index]
        _f.columns = [T("col_count")]
        st.bar_chart(_f, y=T("col_count"))
    else:
        st.caption(T("no_applications"))

    st.markdown(T("source_eff"))
    if not stats["src_df"].empty:
        _src = stats["src_df"].rename(columns={
            "source":      T("col_source"),
            "total":       T("col_total"),
            "grade_a":     T("col_grade_a"),
            "grade_b":     T("col_grade_b"),
            "applied":     T("col_applied"),
            "interviews":  T("col_interviews"),
            "grade_a_pct": T("col_grade_a_pct"),
            "reply_rate":  T("col_reply_rate"),
        })
        st.dataframe(
            _src[[T("col_source"), T("col_total"), T("col_grade_a"), T("col_grade_b"),
                  T("col_applied"), T("col_interviews"), T("col_grade_a_pct"), T("col_reply_rate")]],
            use_container_width=True,
            hide_index=True,
        )
    else:
        st.caption(T("no_data"))

    col_weekly, col_daily = st.columns([3, 2])

    with col_weekly:
        st.markdown(T("weekly_trend"))
        if not stats["week_df"].empty:
            st.bar_chart(stats["week_df"], y="cnt")
        else:
            st.caption(T("no_applications"))

    with col_daily:
        st.markdown(T("daily_trend"))
        if not stats["daily_df"].empty and stats["daily_df"]["cnt"].sum() > 0:
            st.bar_chart(stats["daily_df"], y="cnt")
        else:
            st.caption(T("no_applications"))

with st.expander(T("log_expander"), expanded=False):
    log_path = Path("logs/pipeline.log")
    if log_path.exists():
        _log_text = log_path.read_text(encoding="utf-8", errors="replace")
        _log_lines = _log_text.splitlines()
        st.code("\n".join(_log_lines[-100:]), language=None)
        st.caption(T("log_caption").format(size=log_path.stat().st_size / 1024))
    else:
        st.info(T("no_log"))

    _col1, _col2, _col3 = st.columns(3)
    with _col1:
        if st.button(T("run_btn"), use_container_width=True):
            import subprocess
            import os
            script = str(Path(__file__).parent / "run_pipeline.sh")
            st.info(T("run_starting"))
            subprocess.Popen(
                ["/bin/bash", script],
                cwd=str(Path(__file__).parent),
                stdout=open(log_path, "a"),
                stderr=subprocess.STDOUT,
                start_new_session=True,
            )
            st.success(T("run_started"))
    with _col2:
        if st.button(T("retry_errors_btn"), use_container_width=True):
            import os
            from dotenv import load_dotenv
            from utils.db import init_db, reset_errors_to_unscored
            from phase2_scorer import score_jobs
            load_dotenv()
            _db_path = os.getenv("DB_PATH", "./data/jobs.db")
            _qdrant_path = os.getenv("QDRANT_PATH", "./qdrant_data")
            _conn = init_db(_db_path)
            _n = reset_errors_to_unscored(_conn)
            _conn.close()
            if _n:
                with st.spinner(T("retry_errors_spin")):
                    score_jobs(db_path=_db_path, qdrant_path=_qdrant_path)
                st.cache_data.clear()
                st.rerun()
            else:
                st.info(T("retry_errors_none"))
    with _col3:
        if st.button(T("refresh_btn"), use_container_width=True):
            st.rerun()

st.divider()

# ── Main columns ───────────────────────────────────────────────────────────────

left, right = st.columns([4, 6])

# ════════════════════════════════════════════════════════════════════════════════
# LEFT COLUMN
# ════════════════════════════════════════════════════════════════════════════════

with left:
    # ── Filters ──
    with st.expander(T("filter_expander"), expanded=True):
        fit_grade_filter = st.multiselect(
            T("filter_grade"), ["A", "B", "C"], default=["A", "B"],
            key="filter_grade",
        )
        lang_filter = st.multiselect(
            T("filter_lang"),
            ["en_required", "de_plus", "de_required", "unknown"],
            default=["en_required", "de_plus", "unknown"],
            key="filter_lang",
        )
        source_filter = st.multiselect(
            T("filter_source"),
            ["arbeitnow", "englishjobs", "remotive", "jobicy",
             "relocateme", "bundesagentur", "greenhouse", "lever",
             "wearedevelopers", "heise", "jobware", "ashby", "wttj",
             "weworkremotely", "personio", "germantechjobs",
             "linkedin", "stepstone", "other"],
            default=["arbeitnow", "englishjobs", "remotive", "jobicy",
                     "relocateme", "bundesagentur", "greenhouse", "lever",
                     "wearedevelopers", "heise", "jobware", "ashby", "wttj",
                     "weworkremotely", "personio", "germantechjobs",
                     "linkedin", "stepstone", "other"],
            key="filter_source",
        )
        status_filter = st.multiselect(
            T("filter_status"),
            ["un-scored", "scored", "applied", "interview_1", "interview_2",
             "offer", "rejected", "ghosted", "skipped", "error", "expired"],
            default=["scored"],
            key="filter_status",
        )
        location_filter = st.text_input(
            T("filter_location"),
            placeholder=T("filter_location_ph"),
            key="filter_location",
        )
        remote_filter = st.checkbox(T("filter_remote"), value=False, key="filter_remote")

    # ── Job table ──
    df = fetch_jobs(conn, fit_grade_filter, lang_filter, source_filter, status_filter,
                    location_kw=location_filter, include_remote=remote_filter)

    if df.empty:
        st.info(T("no_jobs"))
        selected_job_id = None
    else:
        event = st.dataframe(
            df[["title", "company", "location", "fit_grade",
                "match_score", "jd_language_req", "source", "status"]],
            use_container_width=True,
            hide_index=True,
            selection_mode="single-row",
            on_select="rerun",
            key="job_table",
        )

        selected_rows = event.selection.get("rows", []) if event.selection else []
        if selected_rows:
            st.session_state["selected_idx"] = selected_rows[0]

        selected_job_id = (
            df.iloc[st.session_state["selected_idx"]]["id"]
            if "selected_idx" in st.session_state
            and st.session_state["selected_idx"] < len(df)
            else None
        )

    # ── LinkedIn search buttons ──
    st.markdown(T("linkedin_search"))
    linkedin_urls = config.get("linkedin_search_urls", [])
    if linkedin_urls:
        link_cols = st.columns(len(linkedin_urls))
        for i, url in enumerate(linkedin_urls):
            label = url.split("keywords=")[1].split("&")[0].replace("+", " ")
            with link_cols[i]:
                st.link_button(f"🔍 {label}", url, use_container_width=True)

    # ── Search targets manager ──
    with st.expander(T("mgr_expander"), expanded=False):
        cfg = load_config()

        tab_kw, tab_gh, tab_lv = st.tabs([T("tab_kw"), T("tab_gh"), T("tab_lv")])

        def _lines(lst: list) -> str:
            return "\n".join(str(x) for x in (lst or []))

        def _parse_lines(text: str) -> list[str]:
            return [ln.strip() for ln in text.splitlines() if ln.strip()]

        with tab_kw:
            st.caption(T("kw_hint"))
            an_kw = st.text_area(
                "Arbeitnow Keywords",
                value=_lines(cfg.get("arbeitnow", {}).get("keywords", [])),
                height=200,
                key="mgr_an_kw",
                label_visibility="collapsed",
            )
            ba_kw = st.text_area(
                T("ba_kw_label"),
                value=_lines(cfg.get("bundesagentur", {}).get("keywords", [])),
                height=160,
                key="mgr_ba_kw",
            )
            if st.button(T("save_kw_btn"), key="save_kw", use_container_width=True):
                cfg["arbeitnow"]["keywords"]    = _parse_lines(an_kw)
                cfg["bundesagentur"]["keywords"] = _parse_lines(ba_kw)
                with open("config/search_targets.yaml", "w", encoding="utf-8") as f:
                    yaml.dump(cfg, f, allow_unicode=True, sort_keys=False, default_flow_style=False)
                load_config.clear()
                st.success(T("save_kw_ok"))

        with tab_gh:
            st.caption(T("gh_hint"))
            gh_slugs = st.text_area(
                T("gh_label"),
                value=_lines(cfg.get("greenhouse", {}).get("companies", [])),
                height=280,
                key="mgr_gh",
                label_visibility="collapsed",
            )
            if st.button(T("save_gh_btn"), key="save_gh", use_container_width=True):
                if "greenhouse" not in cfg:
                    cfg["greenhouse"] = {}
                cfg["greenhouse"]["companies"] = _parse_lines(gh_slugs)
                with open("config/search_targets.yaml", "w", encoding="utf-8") as f:
                    yaml.dump(cfg, f, allow_unicode=True, sort_keys=False, default_flow_style=False)
                load_config.clear()
                st.success(T("save_gh_ok"))

        with tab_lv:
            st.caption(T("lv_hint"))
            lv_slugs = st.text_area(
                T("lv_label"),
                value=_lines(cfg.get("lever", {}).get("companies", [])),
                height=280,
                key="mgr_lv",
                label_visibility="collapsed",
            )
            if st.button(T("save_lv_btn"), key="save_lv", use_container_width=True):
                if "lever" not in cfg:
                    cfg["lever"] = {}
                cfg["lever"]["companies"] = _parse_lines(lv_slugs)
                with open("config/search_targets.yaml", "w", encoding="utf-8") as f:
                    yaml.dump(cfg, f, allow_unicode=True, sort_keys=False, default_flow_style=False)
                load_config.clear()
                st.success(T("save_lv_ok"))

    # ── Manual add ──
    with st.expander(T("manual_expander"), expanded=False):
        with st.form("manual_add"):
            company  = st.text_input(T("company_field"))
            title    = st.text_input(T("title_field"))
            url      = st.text_input(T("url_field"))
            location = st.text_input(T("location_field"))
            source   = st.selectbox(T("source_field"), ["linkedin", "stepstone", "bundesagentur", "other"])
            raw_jd   = st.text_area(T("jd_field"), height=200)
            submit   = st.form_submit_button(T("manual_submit"))

        if submit:
            if not all([company, title, url, raw_jd]):
                st.error(T("manual_required"))
            else:
                job = {
                    "id":          make_id(url),
                    "company":     company,
                    "title":       title,
                    "url":         url,
                    "source":      source,
                    "source_tier": "manual",
                    "location":    location,
                    "raw_jd_text": raw_jd,
                    "fetched_at":  utcnow_iso(),
                    "status":      "un-scored",
                }
                added = upsert_job(conn, job)
                if added:
                    st.success(T("manual_added"))
                else:
                    st.warning(T("manual_exists"))

# ════════════════════════════════════════════════════════════════════════════════
# RIGHT COLUMN
# ════════════════════════════════════════════════════════════════════════════════

with right:
    if "selected_idx" not in st.session_state or not selected_job_id:
        st.info(T("select_prompt"))
    else:
        job = fetch_job_detail(conn, selected_job_id)
        if job is None:
            st.warning(T("not_found"))
        else:
            grade = job.get("fit_grade") or "C"
            st.subheader(f"{GRADE_ICON.get(grade, '')} {job['title']}")
            _translated_flag = T("translated_flag") if job.get("translated_jd_text") else ""
            st.caption(f"{job['company']} · {job.get('location') or '—'} · {job['source']}{_translated_flag}")

            # ── Duplicate application warning ─────────────────────────────
            _prev = get_company_applications(conn, job["company"], job["id"])
            if _prev:
                _status_label = {
                    "applied":     T("status_applied"),
                    "interview_1": T("status_interview_1"),
                    "interview_2": T("status_interview_2"),
                    "offer":       T("status_offer"),
                    "rejected":    T("status_rejected"),
                }
                _prev_lines = "　".join(
                    f"**{p['title']}**（{_status_label.get(p['status'], p['status'])}）"
                    for p in _prev
                )
                st.warning(T("dup_warning").format(company=job["company"], roles=_prev_lines))

            # ── Visa banner + deep analysis ───────────────────────────────
            visa = job.get("visa_restriction") or "unclear"
            _visa_analysis = job.get("visa_analysis")
            _visa_banner_label = {
                "eu_only":   T("visa_eu_only"),
                "sponsored": T("visa_sponsored"),
                "open":      T("visa_open"),
                "unclear":   T("visa_unclear"),
            }.get(visa, visa)
            _visa_expander_label = T("visa_expander_pfx") + _visa_banner_label + ("  ✓" if _visa_analysis else "")

            with st.expander(_visa_expander_label, expanded=(visa == "eu_only" and not _visa_analysis)):
                if _visa_analysis:
                    st.markdown(_visa_analysis)
                else:
                    if visa == "eu_only":
                        st.warning(T("visa_warning_text"))
                    elif visa == "sponsored":
                        st.success(T("visa_sponsored_txt"))
                    else:
                        st.caption(T("visa_unclear_txt"))
                if st.button(
                    T("visa_analyze_btn") if not _visa_analysis else T("visa_reanalyze_btn"),
                    key=f"visa_{job['id']}",
                    type="primary" if (visa == "eu_only" and not _visa_analysis) else "secondary",
                ):
                    import os
                    from dotenv import load_dotenv
                    from utils.visa_checker import analyze_visa_compatibility
                    load_dotenv()
                    with st.spinner(T("visa_analyzing")):
                        analyze_visa_compatibility(
                            job["id"],
                            db_path=os.getenv("DB_PATH", "./data/jobs.db"),
                            lang=_lang(),
                        )
                    st.cache_data.clear()
                    st.rerun()

            c1, c2, c3, c4, c5 = st.columns([0.9, 0.7, 0.9, 1.1, 1.5])
            c1.metric("Match Score",  job.get("match_score") or "—")
            c2.metric("Fit Grade",    grade)
            LANG_LABEL = {
                "en_required": "🇬🇧 EN",
                "de_required": "🇩🇪 DE",
                "de_plus":     "🇩🇪+EN",
                "unknown":     "—",
            }
            c3.metric("Language Req", LANG_LABEL.get(job.get("jd_language_req") or "unknown", job.get("jd_language_req") or "—"))
            VISA_LABEL = {"open": "🟢 open", "eu_only": "🔴 EU only",
                          "sponsored": "🟢 spons.", "unclear": "⚪ unclear"}
            c4.metric("Visa",         VISA_LABEL.get(visa, visa))
            CONTRACT_LABEL = {
                "permanent": T("contract_permanent"),
                "contract":  T("contract_contract"),
                "freelance": T("contract_freelance"),
                "unknown":   T("contract_unknown"),
            }
            c5.metric("Contract",     CONTRACT_LABEL.get(job.get("contract_type") or "unknown", "—"))

            # ── Full JD ──────────────────────────────────────────────────────
            with st.expander(T("jd_expander"), expanded=False):
                _translated = job.get("translated_jd_text")
                if _translated:
                    _jd_tab1, _jd_tab2 = st.tabs([T("jd_tab_en"), T("jd_tab_de")])
                    with _jd_tab1:
                        st.text(_translated)
                    with _jd_tab2:
                        st.text(job.get("raw_jd_text") or "")
                else:
                    st.text(job.get("raw_jd_text") or T("jd_empty"))

            # ── Salary ───────────────────────────────────────────────────────
            _salary_estimate = job.get("salary_estimate")
            _salary_label = T("salary_prefix") + ("  ✓" if _salary_estimate else "")
            if job.get("salary_range"):
                _salary_label += T("salary_jd_shown").format(salary=job["salary_range"])
            with st.expander(_salary_label, expanded=False):
                if _salary_estimate:
                    st.markdown(_salary_estimate)
                else:
                    if job.get("salary_range"):
                        st.caption(T("salary_jd_cap").format(salary=job["salary_range"]))
                    st.caption(T("salary_not_gen"))

                # ── Levels.fyi cache freshness indicator ─────────────────────
                import os
                from dotenv import load_dotenv
                from utils.levels_scraper import cache_info, clear_cache, _role_slug, _location_slug
                load_dotenv()
                _ci = cache_info(job["title"], job.get("location"), job.get("contract_type"))
                if _ci:
                    try:
                        from datetime import datetime, timezone
                        _age_days = (
                            datetime.now(timezone.utc)
                            - datetime.fromisoformat(_ci["fetched_at"])
                        ).days
                    except Exception:
                        _age_days = 0
                    _tpl  = "levels_cache_fresh" if _ci["is_fresh"] else "levels_cache_stale"
                    _locs = " + ".join(_ci["location_slugs"])
                    st.caption(T(_tpl).format(
                        location=_locs,
                        count=_ci["layer_count"],
                        days=_age_days,
                    ))
                else:
                    st.caption(T("levels_cache_none"))

                _s_col1, _s_col2, _s_col3, _s_col4, _s_col5 = st.columns(5)
                with _s_col1:
                    if st.button(
                        T("salary_gen_btn") if not _salary_estimate else T("salary_regen_btn"),
                        key=f"salary_{job['id']}",
                        use_container_width=True,
                        type="primary" if not _salary_estimate else "secondary",
                    ):
                        from utils.salary_estimator import estimate_salary
                        with st.spinner(T("salary_estimating")):
                            estimate_salary(
                                job["id"],
                                db_path=os.getenv("DB_PATH", "./data/jobs.db"),
                                lang=_lang(),
                            )
                        st.cache_data.clear()
                        st.rerun()
                with _s_col2:
                    if st.button(
                        T("levels_refresh_btn"),
                        key=f"levels_refresh_{job['id']}",
                        use_container_width=True,
                    ):
                        _rs = _role_slug(job["title"])
                        _ls = _location_slug(job.get("location"), job.get("contract_type"))
                        from utils.levels_scraper import FALLBACK_CHAIN, fetch_levels_data
                        _chain = FALLBACK_CHAIN.get(_ls, [_ls, "global"])
                        with st.spinner(T("levels_refreshing")):
                            for _slug in _chain:
                                clear_cache(_rs, _slug)
                            fetch_levels_data(job["title"], job.get("location"), job.get("contract_type"))
                        st.rerun()
                with _s_col3:
                    _q = url_quote(f"{job['company']} {job['title']} salary")
                    st.link_button("Glassdoor", f"https://www.glassdoor.com/Search/results.htm?keyword={_q}", use_container_width=True)
                with _s_col4:
                    st.link_button("Kununu", f"https://www.kununu.com/de/search?term={url_quote(job['company'])}", use_container_width=True)
                with _s_col5:
                    _lq = url_quote(job["company"])
                    st.link_button("Levels.fyi", f"https://www.levels.fyi/companies/{_lq}/salaries/", use_container_width=True)

            # ── Company Research ──────────────────────────────────────────────
            _research = job.get("company_research")
            with st.expander(T("research_expander") + ("  ✓" if _research else ""), expanded=False):
                if _research:
                    st.markdown(_research)
                else:
                    st.caption(T("research_not_gen"))
                _r_col1, _r_col2 = st.columns(2)
                with _r_col1:
                    if st.button(
                        T("research_gen_btn") if not _research else T("research_regen_btn"),
                        key=f"research_{job['id']}",
                        use_container_width=True,
                        type="primary" if not _research else "secondary",
                    ):
                        import os
                        from dotenv import load_dotenv
                        from utils.company_researcher import research_company
                        load_dotenv()
                        with st.spinner(T("research_spinner").format(company=job["company"])):
                            research_company(
                                job["id"],
                                db_path=os.getenv("DB_PATH", "./data/jobs.db"),
                                lang=_lang(),
                            )
                        st.cache_data.clear()
                        st.rerun()
                with _r_col2:
                    st.link_button(
                        T("kununu_btn"),
                        f"https://www.kununu.com/de/search?term={url_quote(job['company'])}",
                        use_container_width=True,
                    )

            # Top 3 reasons
            st.markdown(T("reasons_title"))
            reasons_raw = job.get("top_3_reasons")
            reasons = []
            if reasons_raw:
                try:
                    reasons = json.loads(reasons_raw)
                    if not isinstance(reasons, list):
                        reasons = [str(reasons)]
                except json.JSONDecodeError:
                    reasons = [reasons_raw]
            if reasons:
                for r in reasons:
                    st.markdown(f"- {r}")
            else:
                st.caption(T("not_scored_cap"))

            st.divider()

            # Cover letter editor
            st.markdown(T("cl_title"))
            edited_cl = st.text_area(
                label="cover_letter",
                value=job.get("cover_letter_draft") or "",
                height=280,
                label_visibility="collapsed",
                key=f"cl_{job['id']}",
            )
            word_count = len(edited_cl.split()) if edited_cl.strip() else 0
            st.caption(T("cl_word_count").format(n=word_count))
            if word_count > 400:
                st.warning(T("cl_too_long"))

            # Tone regeneration
            _tone_col, _btn_col = st.columns([3, 2])
            with _tone_col:
                _tone = st.selectbox(
                    T("tone_label"),
                    options=["formal", "startup", "concise"],
                    format_func=lambda x: {
                        "formal":  T("tone_formal"),
                        "startup": T("tone_startup"),
                        "concise": T("tone_concise"),
                    }[x],
                    key=f"tone_{job['id']}",
                    label_visibility="collapsed",
                )
            with _btn_col:
                if st.button(T("regen_cl_btn"), use_container_width=True, key=f"regen_cl_{job['id']}"):
                    import os
                    from dotenv import load_dotenv
                    from phase2_scorer import regenerate_cover_letter
                    load_dotenv()
                    with st.spinner(T("regen_cl_spinner").format(tone=_tone)):
                        regenerate_cover_letter(
                            job["id"],
                            tone=_tone,
                            db_path=os.getenv("DB_PATH", "./data/jobs.db"),
                            qdrant_path=os.getenv("QDRANT_PATH", "./qdrant_data"),
                        )
                    st.cache_data.clear()
                    st.rerun()

            if edited_cl.strip():
                import io
                from docx import Document as DocxDocument
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import cm
                from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

                _file_stem = f"cover_letter_{job['company']}_{job['title']}".replace(" ", "_")

                # ── Build .docx ──────────────────────────────────────────────
                _doc = DocxDocument()
                _doc.add_heading(f"{job['title']} @ {job['company']}", level=1)
                for para in edited_cl.strip().split("\n"):
                    _doc.add_paragraph(para)
                _docx_buf = io.BytesIO()
                _doc.save(_docx_buf)

                # ── Build .pdf ───────────────────────────────────────────────
                _pdf_buf = io.BytesIO()
                _pdf_doc = SimpleDocTemplate(
                    _pdf_buf, pagesize=A4,
                    leftMargin=2.5 * cm, rightMargin=2.5 * cm,
                    topMargin=2.5 * cm, bottomMargin=2.5 * cm,
                )
                _styles = getSampleStyleSheet()
                _story = [
                    Paragraph(f"{job['title']} @ {job['company']}", _styles["Heading1"]),
                    Spacer(1, 12),
                ]
                for para in edited_cl.strip().split("\n"):
                    if para.strip():
                        _story.append(Paragraph(para, _styles["Normal"]))
                        _story.append(Spacer(1, 6))
                    else:
                        _story.append(Spacer(1, 10))
                _pdf_doc.build(_story)

                # ── Download buttons (side by side) ──────────────────────────
                _dl_col1, _dl_col2 = st.columns(2)
                with _dl_col1:
                    st.download_button(
                        label=T("download_docx"),
                        data=_docx_buf.getvalue(),
                        file_name=f"{_file_stem}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_docx_{job['id']}",
                        use_container_width=True,
                    )
                with _dl_col2:
                    st.download_button(
                        label=T("download_pdf"),
                        data=_pdf_buf.getvalue(),
                        file_name=f"{_file_stem}.pdf",
                        mime="application/pdf",
                        key=f"dl_pdf_{job['id']}",
                        use_container_width=True,
                    )

            st.divider()

            # ── Action buttons (status-conditional) ──────────────────────────
            cur_status = job.get("status", "")

            def _transition(new_status: str) -> None:
                update_status(conn, job["id"], new_status)
                st.session_state.pop("selected_idx", None)
                st.cache_data.clear()
                st.rerun()

            if cur_status in ("scored", "un-scored", "error"):
                btn_cols = st.columns(5)
                with btn_cols[0]:
                    st.link_button(T("open_job_btn"), job["url"], use_container_width=True)
                with btn_cols[1]:
                    if st.button(T("copy_cl_btn"), use_container_width=True, key=f"copy_{job['id']}"):
                        try:
                            pyperclip.copy(edited_cl)
                            st.success(T("copied_ok"))
                        except Exception:
                            st.info(T("copy_docker_msg"))
                with btn_cols[2]:
                    if st.button(T("apply_btn"), use_container_width=True, type="primary", key=f"apply_{job['id']}"):
                        update_status(conn, job["id"], "applied", applied_at=utcnow_iso())
                        set_follow_up(conn, job["id"],
                            (datetime.now(timezone.utc) + timedelta(days=7)).strftime("%Y-%m-%d"))
                        st.session_state.pop("selected_idx", None)
                        st.cache_data.clear()
                        st.rerun()
                with btn_cols[3]:
                    if st.button(T("skip_btn"), use_container_width=True, key=f"skip_{job['id']}"):
                        _transition("skipped")
                with btn_cols[4]:
                    if st.button(T("rescore_btn"), use_container_width=True, key=f"rescore_{job['id']}"):
                        import os
                        from dotenv import load_dotenv
                        from phase2_scorer import score_single_job
                        load_dotenv()
                        with st.spinner(T("rescore_spinner")):
                            score_single_job(job["id"],
                                db_path=os.getenv("DB_PATH", "./data/jobs.db"),
                                qdrant_path=os.getenv("QDRANT_PATH", "./qdrant_data"))
                        st.cache_data.clear()
                        st.rerun()

            elif cur_status == "applied":
                btn_cols = st.columns(3)
                with btn_cols[0]:
                    st.link_button(T("open_job_btn"), job["url"], use_container_width=True)
                with btn_cols[1]:
                    if st.button(T("iv1_btn"), use_container_width=True, type="primary", key=f"iv1_{job['id']}"):
                        import os
                        from dotenv import load_dotenv
                        from phase2_scorer import generate_brief_for_job
                        load_dotenv()
                        with st.spinner(T("iv1_spinner")):
                            generate_brief_for_job(
                                job["id"],
                                db_path=os.getenv("DB_PATH", "./data/jobs.db"),
                                qdrant_path=os.getenv("QDRANT_PATH", "./qdrant_data"),
                                lang=_lang(),
                            )
                        _transition("interview_1")
                with btn_cols[2]:
                    if st.button(T("reject_btn"), use_container_width=True, key=f"rej_{job['id']}"):
                        _transition("rejected")

            elif cur_status == "interview_1":
                btn_cols = st.columns(3)
                with btn_cols[0]:
                    st.link_button(T("open_job_btn"), job["url"], use_container_width=True)
                with btn_cols[1]:
                    if st.button(T("iv2_btn"), use_container_width=True, type="primary", key=f"iv2_{job['id']}"):
                        _transition("interview_2")
                with btn_cols[2]:
                    if st.button(T("reject_btn"), use_container_width=True, key=f"rej_{job['id']}"):
                        _transition("rejected")

            elif cur_status == "interview_2":
                btn_cols = st.columns(3)
                with btn_cols[0]:
                    st.link_button(T("open_job_btn"), job["url"], use_container_width=True)
                with btn_cols[1]:
                    if st.button(T("offer_btn"), use_container_width=True, type="primary", key=f"offer_{job['id']}"):
                        _transition("offer")
                with btn_cols[2]:
                    if st.button(T("reject_btn"), use_container_width=True, key=f"rej_{job['id']}"):
                        _transition("rejected")

            else:  # offer / rejected / skipped / expired — terminal, open only
                if cur_status == "expired":
                    st.warning(T("expired_warning"))
                st.link_button(T("open_job_btn"), job["url"])

            # ── Interview Brief ──────────────────────────────────────────────
            if cur_status in ("interview_1", "interview_2", "offer"):
                st.divider()
                brief_header, brief_btn = st.columns([5, 1])
                with brief_header:
                    st.markdown(T("brief_title"))
                with brief_btn:
                    if st.button(T("brief_regen_btn"), key=f"regen_brief_{job['id']}", use_container_width=True):
                        import os
                        from dotenv import load_dotenv
                        from phase2_scorer import generate_brief_for_job
                        load_dotenv()
                        with st.spinner(T("brief_gen_spinner")):
                            generate_brief_for_job(
                                job["id"],
                                db_path=os.getenv("DB_PATH", "./data/jobs.db"),
                                qdrant_path=os.getenv("QDRANT_PATH", "./qdrant_data"),
                                lang=_lang(),
                            )
                        st.cache_data.clear()
                        st.rerun()

                brief = job.get("interview_brief")
                if brief:
                    st.markdown(brief)
                    st.download_button(
                        label=T("brief_download"),
                        data=f"# {job['title']} @ {job['company']}\n\n{brief}\n",
                        file_name=f"interview_brief_{job['company']}_{job['title']}.md".replace(" ", "_"),
                        mime="text/markdown",
                        key=f"dl_brief_{job['id']}",
                    )
                else:
                    st.caption(T("brief_not_gen"))

            # ── Interview Records ────────────────────────────────────────────
            if cur_status in ("interview_1", "interview_2", "offer", "rejected"):
                st.divider()
                st.markdown(T("records_title"))

                records = get_interview_records(conn, job["id"])

                ROUND_LABEL = {
                    "interview_1": T("round_iv1"),
                    "interview_2": T("round_iv2"),
                    "other":       T("round_other"),
                }
                FORMAT_LABEL = {
                    "phone":     T("fmt_phone"),
                    "video":     T("fmt_video"),
                    "onsite":    T("fmt_onsite"),
                    "technical": T("fmt_technical"),
                }
                STAR = ["", "★☆☆☆☆", "★★☆☆☆", "★★★☆☆", "★★★★☆", "★★★★★"]

                if records:
                    for rec in records:
                        label = ROUND_LABEL.get(rec["round"], rec["round"])
                        date_str = rec.get("interview_date") or "—"
                        with st.expander(
                            f"{label}　{date_str}　{FORMAT_LABEL.get(rec.get('format',''), rec.get('format',''))}　{STAR[rec.get('self_rating') or 0]}",
                            expanded=False,
                        ):
                            if rec.get("interviewer"):
                                st.caption(f"Interviewer: {rec['interviewer']}")
                            if rec.get("questions"):
                                st.markdown(T("questions_label"))
                                st.markdown(rec["questions"])
                            if rec.get("impressions"):
                                st.markdown(T("impressions_label"))
                                st.markdown(rec["impressions"])
                            if st.button(T("del_record_btn"), key=f"del_rec_{rec['id']}", type="secondary"):
                                delete_interview_record(conn, rec["id"])
                                st.cache_data.clear()
                                st.rerun()

                with st.expander(T("add_record_exp"), expanded=not records):
                    with st.form(key=f"interview_record_form_{job['id']}"):
                        _default_round = cur_status if cur_status in ("interview_1", "interview_2") else "other"
                        _round_options = ["interview_1", "interview_2", "other"]
                        ir_round = st.selectbox(
                            T("round_label"),
                            options=_round_options,
                            format_func=lambda x: ROUND_LABEL.get(x, x),
                            index=_round_options.index(_default_round),
                        )
                        ir_col1, ir_col2 = st.columns(2)
                        with ir_col1:
                            ir_date = st.date_input(T("ir_date_label"), value=datetime.now(timezone.utc).date())
                        with ir_col2:
                            ir_format = st.selectbox(
                                T("ir_format_label"),
                                options=["video", "phone", "onsite", "technical"],
                                format_func=lambda x: FORMAT_LABEL.get(x, x),
                            )
                        ir_interviewer = st.text_input(T("interviewer_label"), placeholder="e.g. Sarah, Engineering Manager")
                        ir_questions = st.text_area(T("questions_field"), height=120, placeholder="- Tell me about yourself\n- How do you handle …")
                        ir_rating = st.slider(T("rating_label"), min_value=1, max_value=5, value=3, help=T("rating_help"))
                        ir_impressions = st.text_area(T("impressions_field"), height=80)
                        if st.form_submit_button(T("save_record_btn"), use_container_width=True, type="primary"):
                            add_interview_record(conn, {
                                "job_id":         job["id"],
                                "round":          ir_round,
                                "interview_date": ir_date.isoformat(),
                                "interviewer":    ir_interviewer or None,
                                "format":         ir_format,
                                "questions":      ir_questions or None,
                                "self_rating":    ir_rating,
                                "impressions":    ir_impressions or None,
                                "created_at":     utcnow_iso(),
                            })
                            st.cache_data.clear()
                            st.rerun()

            # ── Notes ────────────────────────────────────────────────────────
            st.divider()
            st.markdown(T("notes_title"))
            note_val = st.text_area(
                label="notes",
                value=job.get("notes") or "",
                height=80,
                placeholder=T("notes_placeholder"),
                label_visibility="collapsed",
                key=f"notes_{job['id']}",
            )
            if st.button(T("save_notes_btn"), key=f"save_notes_{job['id']}"):
                set_notes(conn, job["id"], note_val)
                st.cache_data.clear()
                st.success(T("notes_saved"))

            # ── Follow-up section (all active pipeline stages) ───────────────
            if cur_status in PIPELINE_STATUSES:
                st.divider()
                st.markdown(T("followup_title"))

                current_date = job.get("follow_up_at")
                default_date = (
                    datetime.fromisoformat(current_date).date()
                    if current_date
                    else (datetime.now(timezone.utc) + timedelta(days=7)).date()
                )

                fu_col1, fu_col2 = st.columns([3, 1])
                with fu_col1:
                    new_date = st.date_input(
                        T("followup_title"),
                        value=default_date,
                        key=f"fu_date_{job['id']}",
                        label_visibility="collapsed",
                    )
                with fu_col2:
                    if st.button(T("followup_save_btn"), key=f"fu_save_{job['id']}", use_container_width=True):
                        set_follow_up(conn, job["id"], new_date.isoformat())
                        st.cache_data.clear()
                        st.success(f"{new_date.isoformat()}")

                if current_date:
                    today = today_iso()
                    if current_date <= today:
                        st.warning(T("followup_overdue").format(date=current_date))
                    else:
                        days_left = (
                            datetime.fromisoformat(current_date).date()
                            - datetime.now(timezone.utc).date()
                        ).days
                        st.caption(T("followup_days_left").format(n=days_left, date=current_date))
