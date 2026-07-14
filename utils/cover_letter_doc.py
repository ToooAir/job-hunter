"""Render a cover-letter text into downloadable .pdf / .docx bytes.

Shared by the dashboard (phase3_dashboard) and the Apply Review queue so both
produce byte-identical documents from the same letter. Forms that take the
cover letter as a file upload (not a paste-able textarea) need this.

reportlab's Paragraph parses its input as mini-XML, so bare "&"/"<"/">" —
common in German company names like "Merz Pharma & Co. KG" — must be escaped
or the build raises a paraparser syntax error. python-docx takes text
literally, so only the PDF path escapes.
"""
from __future__ import annotations

import io
import re
from xml.sax.saxutils import escape


def file_stem(company: str, title: str) -> str:
    """Filesystem-safe basename (no extension) for a downloaded letter."""
    stem = f"cover_letter_{company}_{title}"
    return re.sub(r"[^\w.-]+", "_", stem).strip("_") or "cover_letter"


def build_pdf(text: str, title: str, company: str) -> bytes:
    """A4 PDF: "<title> @ <company>" heading, one paragraph per source line."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    story = [
        Paragraph(escape(f"{title} @ {company}"), styles["Heading1"]),
        Spacer(1, 12),
    ]
    for para in text.strip().split("\n"):
        if para.strip():
            story.append(Paragraph(escape(para), styles["Normal"]))
            story.append(Spacer(1, 6))
        else:
            story.append(Spacer(1, 10))
    doc.build(story)
    return buf.getvalue()


def build_docx(text: str, title: str, company: str) -> bytes:
    """A .docx with a "<title> @ <company>" heading and one paragraph per line."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading(f"{title} @ {company}", level=1)
    for para in text.strip().split("\n"):
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
